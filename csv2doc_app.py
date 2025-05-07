# ─────────────────────────────────────────────────────────────────────────────
# CSV to Document Converter
# ─────────────────────────────────────────────────────────────────────────────
# This script allows users to upload CSV files and convert them into different document formats (CSV, HTML, DOCX).
# It also provides options for grouping, sorting, and filtering the data before conversion.
# In the process of developing a site for documentation, via GitHub Pages,
# I wanted to create a simple and user-friendly interface for converting CSV files into various document formats.
# So developed this script to help users easily manage and convert their CSV data into more readable formats.
# It uses Streamlit for the web interface and Pandas for data manipulation.

import traceback
import streamlit as st
import pandas as pd
import io
import base64
import json
import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from pathlib import Path

# ---------------------------------------------------------------------------
# Image handling utilities
# ---------------------------------------------------------------------------

BASE_DIR   = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR / "assets"

def get_float_image_path() -> str:
    """
    Returns an absolute path to “float_image.png” located inside ./assets.
    If the file does not exist, a harmless online placeholder is returned so
    that the UI never shows a broken-image icon.
    """
    img_path = ASSETS_DIR / "float_image.png"

    if img_path.exists():
        return str(img_path)

    # Fallback (any publicly available placeholder will do)
    return "https://via.placeholder.com/150"


# Set page config - must be the first Streamlit command
st.set_page_config(page_title="CSV to Document Converter", layout="wide", page_icon="favicon.png")

# Replace the former hard-coded value
# float_image_path = get_float_image_path()
# st.info(f"BASE_DIR = {BASE_DIR}, ASSETS_DIR = {ASSETS_DIR}, float_image_path = {float_image_path}")
float_image_path = "csv2doc.png" 

# Create a data directory if it doesn't exist
def ensure_data_dir_exists():
    os.makedirs("data", exist_ok=True)
    os.makedirs("data/users", exist_ok=True)
    os.makedirs("data/csv_cache", exist_ok=True)

# Ensure the data directory exists
ensure_data_dir_exists()
query_params = st.query_params

# User management functions
def get_user_id():
    """Get a unique user ID or create one if it doesn't exist"""
    # Check if user_id exists in session state
    if "user_id" not in st.session_state:
        # Try to get user_id from query parameters (cookie-like behavior)
        if "user_id" in query_params and query_params["user_id"]:
            # Use the user_id from query parameters
            st.session_state.user_id = query_params["user_id"][0]
        elif "persistent_user_id" in st.session_state:
            # Use previously stored ID
            st.session_state.user_id = st.session_state.persistent_user_id
        else:
            # Generate a new user ID
            st.session_state.user_id = str(uuid.uuid4())
            st.session_state.persistent_user_id = st.session_state.user_id
            # Store user_id in query parameters (cookie-like behavior)
            query_params["user_id"] = st.session_state.user_id

    # Ensure user_id is always in query parameters - check existence first
    if "user_id" not in query_params or not query_params["user_id"]:
        query_params["user_id"] = st.session_state.user_id

    return st.session_state.user_id

def get_user_history_path(user_id):
    """Get path to user's history file"""
    return f"data/users/{user_id}/history.json"

def save_to_user_history(user_id, file_name, cache_id=None):
    """
    Save file information to user history
    If cache_id is provided, it will be used to link to the cached file
    """
    # Choose an appropriate path for user history
    history_path = get_user_history_path(user_id)
    os.makedirs(os.path.dirname(history_path), exist_ok=True)
    user_history = []

    # Load existing history if it exists
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            user_history = json.load(f)

    # Check if this file is already in history
    file_exists = False
    for entry in user_history:
        if isinstance(entry, dict) and entry.get("file_name") == file_name:
            # Update the existing entry with new timestamp and cache_id
            entry["timestamp"] = datetime.now().isoformat()
            if cache_id:
                entry["cache_id"] = cache_id
            file_exists = True
            break

    # If file doesn't exist in history, add it
    if not file_exists:
        csv_data = {
            "file_name": file_name,
            "timestamp": datetime.now().isoformat()
        }
        if cache_id:
            csv_data["cache_id"] = cache_id
        user_history.append(csv_data)

    # Save updated history
    with open(history_path, 'w', encoding='utf-8') as f:
        json.dump(user_history, f)

def get_user_history(user_id):
    history_path = get_user_history_path(user_id)
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

# CSV cache management
def ensure_csv_cache_dir_exists():
    """Create the CSV cache directory if it doesn't exist"""
    os.makedirs("data/csv_cache", exist_ok=True)

def get_file_hash(df, file_name):
    """Generate a hash for the DataFrame content and file name to use as identifier"""
    # Convert DataFrame to CSV string and combine with file name before hashing
    csv_string = df.to_csv(index=False)
    # Combine file name and content to ensure uniqueness even for identical content
    combined_string = f"{file_name}_{csv_string}"
    return str(hash(combined_string))

def cache_csv(df, file_name, options=None):
    """
    Save DataFrame to session cache and disk cache with its file name and options
    Returns a cache_id that can be used to retrieve the DataFrame later

    Parameters:
    -----------
    df : pd.DataFrame
        The DataFrame to cache
    file_name : str
        The name of the file
    options : dict, optional
        Dictionary containing grouping, sorting, and filtering options
    """
    # Ensure cache directory exists
    ensure_csv_cache_dir_exists()

    # Generate a hash for the file content and name
    file_hash = get_file_hash(df, file_name)

    # Check if we already have this file in the cache
    cache_path = f"data/csv_cache/{file_hash}.csv"

    # Get current user ID for user-specific options
    user_id = get_user_id()

    # Create user-specific options directory if it doesn't exist
    user_options_dir = f"data/users/{user_id}/options"
    os.makedirs(user_options_dir, exist_ok=True)

    # Define paths for global and user-specific options
    global_options_path = f"data/csv_cache/{file_hash}_options.json"
    user_options_path = f"{user_options_dir}/{file_hash}_options.json"

    # Save to disk cache if not already there
    if not os.path.exists(cache_path):
        df.to_csv(cache_path, index=False)

    # Save options to user-specific path if provided
    if options:
        with open(user_options_path, 'w', encoding='utf-8') as f:
            json.dump(options, f)

    # Update the uploaded_files_history.json
    history_path = "uploaded_files_history.json"
    history = {}
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                history = json.load(f)
            except json.JSONDecodeError:
                history = {}

    # Store file info in history
    history[file_hash] = {
        "file_name": file_name,
        "cache_path": cache_path,
        "options_path": global_options_path,  # Keep this for backward compatibility
        "timestamp": datetime.now().isoformat()
    }

    # Also store the file name to options mapping
    if "file_options" not in history:
        history["file_options"] = {}

    # We no longer update global options mapping
    # This is now handled per user

    with open(history_path, 'w', encoding='utf-8') as f:
        json.dump(history, f)

    # Also store in session cache for current session
    if "csv_cache" not in st.session_state:
        st.session_state.csv_cache = {}

    st.session_state.csv_cache[file_hash] = {
        "df": df,
        "file_name": file_name,
        "options": options,
        "timestamp": datetime.now().isoformat()
    }

    return file_hash

def get_cached_csv(cache_id):
    """
    Retrieve DataFrame and options from cache (session or disk)
    Returns a tuple (DataFrame, options) or (DataFrame, default_options) if options not found
    Returns (None, None) if DataFrame not found
    """
    # First try session cache
    if "csv_cache" in st.session_state and cache_id in st.session_state.csv_cache:
        cache_entry = st.session_state.csv_cache[cache_id]
        return cache_entry["df"], cache_entry.get("options")

    # If not in session, try disk cache
    cache_path = f"data/csv_cache/{cache_id}.csv"

    # Get current user ID for user-specific options
    user_id = get_user_id()

    # Define paths for global and user-specific options
    global_options_path = f"data/csv_cache/{cache_id}_options.json"
    user_options_path = f"data/users/{user_id}/options/{cache_id}_options.json"

    if os.path.exists(cache_path):
        try:
            df = pd.read_csv(cache_path, dtype=str)

            # Get file name from history
            file_name = get_file_name_from_history(cache_id)

            # Try to load options if they exist - first check user-specific options
            options = None

            # First try user-specific options
            if os.path.exists(user_options_path):
                try:
                    with open(user_options_path, 'r', encoding='utf-8') as f:
                        options = json.load(f)
                except Exception as e:
                    st.warning(f"Error loading user options: {e}")

            # If no user-specific options, try global options (for backward compatibility)
            if options is None and os.path.exists(global_options_path):
                try:
                    with open(global_options_path, 'r', encoding='utf-8') as f:
                        options = json.load(f)
                except Exception as e:
                    st.warning(f"Error loading global options: {e}")

            # If options still not found, try to get them from file name mapping
            if options is None:
                # Check if we have options for this file name in history
                history_path = "uploaded_files_history.json"
                if os.path.exists(history_path):
                    with open(history_path, 'r', encoding='utf-8') as f:
                        try:
                            history = json.load(f)
                            if "file_options" in history and file_name in history["file_options"]:
                                options = history["file_options"][file_name]
                        except json.JSONDecodeError:
                            pass

            # If still no options, use default options
            if options is None:
                options = {
                    "grouping": {
                        "enabled": False,
                        "columns": []
                    },
                    "sorting": {
                        "enabled": False,
                        "columns": [],
                        "order": "Ascending"
                    },
                    "filtering": {
                        "enabled": False,
                        "filter_columns": [],
                        "filter_conditions": {},
                        "column_filter_type": "None",
                        "include_columns": [],
                        "exclude_columns": []
                    }
                }

            # Add to session cache for future use
            if "csv_cache" not in st.session_state:
                st.session_state.csv_cache = {}

            st.session_state.csv_cache[cache_id] = {
                "df": df,
                "file_name": file_name,
                "options": options,
                "timestamp": datetime.now().isoformat()
            }

            return df, options
        except Exception as e:
            st.error(f"Error loading cached CSV: {e}")

    return None, None

def get_file_name_from_history(cache_id):
    """Get the original file name from the history"""
    history_path = "uploaded_files_history.json"
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                history = json.load(f)
                if cache_id in history:
                    return history[cache_id]["file_name"]
            except json.JSONDecodeError:
                pass

    return f"cached_file_{cache_id}.csv"

def get_all_cached_csvs():
    """Get all cached CSVs for the current session"""
    if "csv_cache" not in st.session_state:
        st.session_state.csv_cache = {}
    return st.session_state.csv_cache

def get_all_disk_cached_csvs():
    """Get all CSVs from the disk cache"""
    history_path = "uploaded_files_history.json"
    if os.path.exists(history_path):
        with open(history_path, 'r', encoding='utf-8') as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                pass

    return {}

# Your existing conversion functions
def convert_df_to_csv(df):
    """Convert DataFrame to CSV string"""
    return df.to_csv(index=False).encode('utf-8')

# python
def convert_df_to_html(df, doc_title=None):
    """Convert DataFrame to HTML string with an optional title"""
    html_table = df.to_html(index=False, border=0)
    style = """
    <style>
      table, th, td {
        border: 1px solid black;
        border-collapse: collapse;
        padding: 5px 10px;
        margin: 5px;
        font-size: 12px;
        font-family: Arial, sans-serif;
      }
    </style>
    """
    title_html = f"<h1>{doc_title}</h1>" if doc_title else ""
    return style + title_html + html_table

def convert_df_to_grouped_html(df: pd.DataFrame, group_cols: list[str], doc_title: str) -> str:
    """
    Return an HTML string in which *df* is split by the given columns.
    If *group_cols* is empty, the plain table is returned.
    """
    if not group_cols:
        return convert_df_to_html(df)

    html_chunks: list[str] = []
    grouped = df.groupby(group_cols)

    if doc_title:
        title = f"<h3>{doc_title}</h3>"
    else:
        title = ", ".join(f"{col} = {val}" for col, val in group_cols.items())
    previous_group = None
    for keys, group in grouped:
        if previous_group is not None and previous_group != group:
            previous_group = group
            # Add a blank line between groups
            html_chunks.append("<br>")
            html_chunks.append(f"<h3>{title}</h3>")
        # Normalise keys → always a tuple so we can zip()
        keys = (keys,) if not isinstance(keys, tuple) else keys
        html_chunks.append(group.to_html(index=False))

    return "\n".join(html_chunks)


# Refactored DOCX conversion function with better error handling
def convert_df_to_docx(df, title="Data Document"):
    """Convert DataFrame to DOCX and return bytes"""
    try:
        doc = Document()
        doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add table
        cols = list(df.columns)
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(cols):
            run = hdr_cells[i].paragraphs[0].add_run(str(col))
            run.bold = True
            run.font.size = Pt(12)

        # Data rows
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                run = cells[i].paragraphs[0].add_run(str(row[col]))
                run.font.size = Pt(10)

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Error creating DOCX file: {str(e)}")
        return None

def convert_df_to_grouped_docx(df: pd.DataFrame, group_cols: list[str], doc_title: str = "") -> bytes:
    """
    Build a DOCX file where *df* is split by *group_cols*.
    Returns the binary content (bytes).
    """
    try:
        from docx import Document
        from docx.shared import Pt

        # Create the document
        doc = Document()
        if doc_title:
            doc.add_heading(doc_title, level=1)

        # If no grouping specified, just convert the whole DataFrame
        if not group_cols:
            return convert_df_to_docx(df, doc_title)

        grouped = df.groupby(group_cols)

        # Loop through groups
        for keys, group in grouped:
            # Ensure `keys` is a tuple (even for a single grouping column)
            keys = (keys,) if not isinstance(keys, tuple) else keys

            # Create a heading for the group based on the keys and group columns
            heading = ", ".join(f"{col} = {val}" for col, val in zip(group_cols, keys))
            doc.add_heading(heading, level=3)

            # Add the table for this group
            table = doc.add_table(rows=1, cols=len(group.columns))
            table.style = "Table Grid"

            # Add Header row
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(group.columns):
                hdr_cells[idx].text = str(col)
                for paragraph in hdr_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            # Add Data rows
            for _, row in group.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            doc.add_paragraph()  # Blank line between groups

        # Serialize the document to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        st.error(f"Error creating grouped DOCX file: {str(e)}")
        return None

def get_download_link(data, filename, text):
    """Generate a download link for the data"""
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:file/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

# Assume: df is your current DataFrame
# download_format: "csv", "html", or "docx"

def get_data_preview_and_download_options(df, download_format):
    # Generate preview HTML based on the format
    if download_format == "csv":
        preview = df.head().to_csv()
        file_content = convert_df_to_csv(df)
        file_ext = "csv"
    elif download_format == "html":
        preview = df.head().to_html()
        file_content = convert_df_to_html(df)
        file_ext = "html"
    elif download_format == "docx":
        preview = "Preview not available for DOCX format." # Or show a summary
        file_content = convert_df_to_docx(df)
        file_ext = "docx"
    else:
        preview = ""
        file_content = None
        file_ext = ""

    # Always return a single preview + download section
    download_link = get_download_link(file_content, f"export.{file_ext}", f"Download File - Table")
    return f"""
        <div id='data-preview'>
            <h4>Preview:</h4>
            <pre>{preview}</pre>
            <h4>Download:</h4>
            {download_link}
        </div>
    """

# In your route or callback or rendering logic:
# Then, in the template/UI rendering, always inject/replace the data_preview_section

# ... keep your existing conversion functions unchanged ...

# Initialize session state for storing dataframes if not exists
if "dataframes" not in st.session_state:
    st.session_state.dataframes = []

# Get the current user ID
user_id = get_user_id()

col_img, col_body = st.columns([1, 10])      # adjust ratios to your liking
if float_image_path and os.path.exists(float_image_path):
    col_img.image(float_image_path, use_container_width=False, width=150)
else:
    st.warning(f"Image not found at {float_image_path}, using placeholder.")
# col_img.image(float_image_path, use_column_width=True, caption="CSV to Document Converter", width=150)

with col_body:
    # st.title("CSV to Document Converter")

    # File uploader
    col1, col2 = st.columns([3, 1])
    with col1:
        uploaded_file = st.file_uploader("Choose a CSV file", type="csv", key="upload")
    with col2:
        st.write("<br>", unsafe_allow_html=True)
        load_example = st.button("Load Example: Emojis.csv 😃")

    # Handle example file loading
    if load_example:
        # Load the emojis.csv file
        example_file_path = "emojis.csv"
        if os.path.exists(example_file_path):
            df = pd.read_csv(example_file_path, dtype=str)

            # Add to dataframes list
            st.session_state.dataframes.append(df)

            # Initialize default options for the new file
            default_options = {
                "grouping": {
                    "enabled": False,
                    "columns": []
                },
                "sorting": {
                    "enabled": False,
                    "columns": [],
                    "order": "Ascending"
                },
                "filtering": {
                    "enabled": False,
                    "filter_columns": [],
                    "filter_conditions": {},
                    "column_filter_type": "None",
                    "include_columns": [],
                    "exclude_columns": []
                }
            }

            # Cache the CSV with default options and get the cache_id
            cache_id = cache_csv(df, "emojis.csv", default_options)

            # Add to user history with the cache_id
            save_to_user_history(user_id, "emojis.csv", cache_id)

            # Store the options in session state for this dataframe
            if "dataframe_options" not in st.session_state:
                st.session_state.dataframe_options = {}
            # The index will be the length of dataframes list - 1
            df_idx = len(st.session_state.dataframes) - 1
            st.session_state.dataframe_options[df_idx] = default_options

            st.success(f"Successfully loaded emojis.csv with {df.shape[0]} rows and {df.shape[1]} columns.")
        else:
            st.error(f"Example file not found at {example_file_path}")

    if uploaded_file is not None:
        try:
            df = pd.read_csv(uploaded_file, dtype=str)

            # Add to dataframes list
            st.session_state.dataframes.append(df)

            # Initialize default options for the new file
            default_options = {
                "grouping": {
                    "enabled": False,
                    "columns": []
                },
                "sorting": {
                    "enabled": False,
                    "columns": [],
                    "order": "Ascending"
                },
                "filtering": {
                    "enabled": False,
                    "filter_columns": [],
                    "filter_conditions": {},
                    "column_filter_type": "None",
                    "include_columns": [],
                    "exclude_columns": []
                }
            }

            # Cache the CSV with default options and get the cache_id
            cache_id = cache_csv(df, uploaded_file.name, default_options)

            # Add to user history with the cache_id
            save_to_user_history(user_id, uploaded_file.name, cache_id)

            # Store the options in session state for this dataframe
            if "dataframe_options" not in st.session_state:
                st.session_state.dataframe_options = {}
            # The index will be the length of dataframes list - 1
            df_idx = len(st.session_state.dataframes) - 1
            st.session_state.dataframe_options[df_idx] = default_options

            st.success(f"Successfully loaded CSV with {df.shape[0]} rows and {df.shape[1]} columns.")
            # download_format = st.selectbox("Format", ["csv", "html", "docx"])
            # enable_grouping = st.checkbox("Group by column")
            # if enable_grouping:
            #     # switch to a single‐select so df[group_col] is a Series
            #     group_col = st.selectbox("Select column to group by", df.columns)
            #     unique_vals = df[group_col].unique().tolist()
            #     # let user pick exactly one group
            #     chosen = st.selectbox("Which group to preview?", unique_vals)
            #     subset = df[df[group_col] == chosen]
            #     st.markdown(
            #         get_data_preview_and_download_options(subset, download_format),
            #         unsafe_allow_html=True
            #     )
            # else:
            #     st.markdown(
            #         get_data_preview_and_download_options(df, download_format),
            #         unsafe_allow_html=True
            #     )

        except Exception as e:
            st.error(f"Error processing the file")
            st.error(f"Exception occurred")
            st.markdown(f"**Exception:**<br>{traceback.format_exc().replace('\n', '<br>')}", unsafe_allow_html=True)

    # Display all loaded CSVs in separate tables
    for idx, df in enumerate(st.session_state.dataframes):
        with st.expander(f"Data Table", expanded=True):
            st.subheader(f"Data Preview - Table")
            st.text(f"Loaded {df.shape[0]} rows and {df.shape[1]} columns.")
            st.dataframe(df.head(10))

            # Document title input per dataframe
            doc_title = st.text_input(f"Document Title for Table",
                                     f"Data Document {idx+1}",
                                     key=f"title_{idx}")

            # Grouping, sorting, and filtering in tabs
            tab1, tab2, tab3 = st.tabs(["Grouping", "Sorting", "Filtering"])

            # Check if we have saved options for this dataframe
            saved_options = None
            if "dataframe_options" in st.session_state and idx in st.session_state.dataframe_options:
                saved_options = st.session_state.dataframe_options[idx]

            with tab1:
                # Use saved grouping options if available
                default_enable_grouping = False
                default_group_cols = []
                if saved_options and "grouping" in saved_options:
                    default_enable_grouping = saved_options["grouping"].get("enabled", False)
                    default_group_cols = saved_options["grouping"].get("columns", [])

                enable_grouping = st.checkbox("Enable grouping by column", 
                                             value=default_enable_grouping, 
                                             key=f"grouping_{idx}")
                group_cols = []
                if enable_grouping and len(df.columns) > 0:
                    group_cols = st.multiselect("Select column(s) to group by", 
                                               df.columns, 
                                               default=default_group_cols,
                                               key=f"group_col_{idx}")

            with tab2:
                # Use saved sorting options if available
                default_enable_sorting = False
                default_sort_cols = []
                default_sort_order = "Ascending"
                if saved_options and "sorting" in saved_options:
                    default_enable_sorting = saved_options["sorting"].get("enabled", False)
                    default_sort_cols = saved_options["sorting"].get("columns", [])
                    default_sort_order = saved_options["sorting"].get("order", "Ascending")

                enable_sorting = st.checkbox("Enable sorting", 
                                           value=default_enable_sorting, 
                                           key=f"sort_enable_{idx}")
                sort_cols = []
                ascending = True
                if enable_sorting:
                    sort_cols = st.multiselect("Select column(s) to sort by",
                                            options=df.columns,
                                            default=default_sort_cols,
                                            key=f"sort_cols_{idx}")
                    sort_order = st.selectbox("Sort order", 
                                            ("Ascending", "Descending"), 
                                            index=0 if default_sort_order == "Ascending" else 1,
                                            key=f"sort_order_{idx}")
                    ascending = sort_order == "Ascending"

            with tab3:
                # Use saved filtering options if available
                default_enable_filtering = False
                default_filter_cols = []
                default_filter_conditions = {}
                default_column_filter_type = "None"
                default_include_columns = []
                default_exclude_columns = []

                if saved_options and "filtering" in saved_options:
                    default_enable_filtering = saved_options["filtering"].get("enabled", False)
                    default_filter_cols = saved_options["filtering"].get("filter_columns", [])
                    default_filter_conditions = saved_options["filtering"].get("filter_conditions", {})
                    default_column_filter_type = saved_options["filtering"].get("column_filter_type", "None")
                    default_include_columns = saved_options["filtering"].get("include_columns", [])
                    default_exclude_columns = saved_options["filtering"].get("exclude_columns", [])

                enable_filtering = st.checkbox("Enable filtering", 
                                             value=default_enable_filtering, 
                                             key=f"filter_enable_{idx}")
                filter_conditions = {}
                include_columns = []
                exclude_columns = []

                if enable_filtering:
                    # Value filtering
                    st.subheader("Filter by values")
                    filter_cols = st.multiselect("Select column(s) to filter by value",
                                              options=df.columns,
                                              default=default_filter_cols,
                                              key=f"filter_cols_{idx}")
                    for col in filter_cols:
                        unique_values = df[col].dropna().unique().tolist()
                        # Get default selected values for this column if available
                        default_values = default_filter_conditions.get(col, [])
                        # Filter default values to ensure they exist in the current unique values
                        valid_default_values = [v for v in default_values if v in unique_values]

                        selected_values = st.multiselect(f"Select values for {col}",
                                                     options=unique_values,
                                                     default=valid_default_values,
                                                     key=f"filter_values_{idx}_{col}")
                        if selected_values:
                            filter_conditions[col] = selected_values

                    # Column filtering
                    st.subheader("Filter by columns")
                    column_filter_type = st.radio(
                        "Column filtering type",
                        ["None", "Include only", "Exclude only"],
                        index=["None", "Include only", "Exclude only"].index(default_column_filter_type),
                        key=f"column_filter_type_{idx}"
                    )

                    if column_filter_type == "Include only":
                        include_columns = st.multiselect(
                            "Select columns to include in the output",
                            options=df.columns,
                            default=default_include_columns,
                            key=f"include_columns_{idx}"
                        )
                    elif column_filter_type == "Exclude only":
                        exclude_columns = st.multiselect(
                            "Select columns to exclude from the output",
                            options=df.columns,
                            default=default_exclude_columns,
                            key=f"exclude_columns_{idx}"
                        )

            # Apply transformations to create the processed dataframe
            processed_df = df.copy()

            # Apply value filtering
            for col, values in filter_conditions.items():
                processed_df = processed_df[processed_df[col].isin(values)]

            # Apply column filtering (include or exclude)
            if include_columns:
                processed_df = processed_df[include_columns]
            elif exclude_columns:
                processed_df = processed_df.drop(columns=exclude_columns)

            # Apply sorting
            if enable_sorting and sort_cols:
                # Filter sort_cols to only include columns that exist in the DataFrame
                valid_sort_cols = [col for col in sort_cols if col in processed_df.columns]
                if valid_sort_cols:
                    processed_df = processed_df.sort_values(by=valid_sort_cols, ascending=ascending)
                else:
                    st.warning(f"None of the selected sort columns exist in the filtered DataFrame. Sorting skipped.")

            # Collect current options
            current_options = {
                "grouping": {
                    "enabled": enable_grouping,
                    "columns": group_cols
                },
                "sorting": {
                    "enabled": enable_sorting,
                    "columns": sort_cols,
                    "order": sort_order if enable_sorting else "Ascending"
                },
                "filtering": {
                    "enabled": enable_filtering,
                    "filter_columns": filter_cols if enable_filtering and 'filter_cols' in locals() else [],
                    "filter_conditions": filter_conditions,
                    "column_filter_type": column_filter_type if enable_filtering else "None",
                    "include_columns": include_columns,
                    "exclude_columns": exclude_columns
                }
            }

            # Save options button
            if st.button("Save current options", key=f"save_options_{idx}"):
                # Get the cache_id for this dataframe
                file_hash = None
                for cache_id, cache_info in st.session_state.csv_cache.items():
                    if cache_info["df"].equals(df):
                        file_hash = cache_id
                        break

                if file_hash:
                    # Update the cache with the current options
                    cache_csv(df, st.session_state.csv_cache[file_hash]["file_name"], current_options)

                    # Update the options in session state
                    if "dataframe_options" not in st.session_state:
                        st.session_state.dataframe_options = {}
                    st.session_state.dataframe_options[idx] = current_options

                    st.success("Options saved successfully!")
                else:
                    st.error("Could not find cache entry for this dataframe")

            # Download options
            st.subheader(f"Download Options for Table")
            download_format = st.radio("Select download format",
                                     ["CSV", "HTML", "DOCX"],
                                     key=f"download_format_{idx}")

            if st.button(f"Generate Document for Table"):
                if download_format == "CSV":
                    csv_data = convert_df_to_csv(processed_df)
                    st.download_button(
                        label=f"Download CSV - Table",
                        data=csv_data,
                        file_name=f"data_export_{idx+1}.csv",
                        mime="text/csv",
                        on_click="ignore"
                    )
                elif download_format == "HTML":
                    if enable_grouping and group_cols:
                        html_data = convert_df_to_grouped_html(processed_df, group_cols, doc_title)
                        html_bytes = html_data.encode('utf-8')  # Convert string to bytes
                    else:
                        html_data = convert_df_to_html(processed_df, doc_title)
                        html_bytes = html_data.encode('utf-8') if isinstance(html_data, str) else html_data

                    st.download_button(
                        label=f"Download HTML - Table",
                        data=html_bytes,
                        file_name=f"grouped_data_{idx+1}.html" if enable_grouping and group_cols else f"data_{idx+1}.html",
                        mime="text/html",
                        on_click="ignore"
                    )
                elif download_format == "DOCX":
                    # Similar handling for DOCX format
                    if enable_grouping and group_cols:
                        docx_data = convert_df_to_grouped_docx(processed_df, group_cols, doc_title)
                    else:
                        docx_data = convert_df_to_docx(processed_df, doc_title)
                    st.download_button(
                        label=f"Download DOCX - Table",
                        data=docx_data,
                        file_name=f"grouped_data_{idx+1}.docx" if enable_grouping and group_cols else f"data_{idx+1}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        on_click="ignore"
                    )
        break

    # Footer
    st.markdown("---")
    st.markdown("CSV to Document Converter | Created with Streamlit")
    st.markdown("This app allows you to upload a CSV file and convert it to different document formats. You can also group the data by a column to create separate tables for each unique value.")
    st.markdown("**Note:** Make sure to upload a valid CSV file.")
    st.markdown("**Disclaimer:** This app is for educational purposes only. Please ensure you have the right to use any data you upload.")
    # ─────────────────────────────────────────────────────────────────────────────
    #  CONVERT A DATAFRAME TO GROUPED HTML
    # ─────────────────────────────────────────────────────────────────────────────
    def convert_df_to_grouped_html(
        df: pd.DataFrame,
        group_cols: list[str],
        doc_title: str | None = None,        # ← NEW, optional third parameter
    ) -> str:
        """
        Convert a DataFrame to an HTML string, grouped by the specified columns.

        Parameters
        ----------
        df         : pd.DataFrame
            The data that will be rendered.
        group_cols : list[str]
            Column names to group by.
        doc_title  : str | None, optional
            An optional document title.  If supplied, a corresponding
            <h1> header is inserted at the top of the generated HTML.

        Returns
        -------
        str
            The rendered HTML string.
        """
        html_parts: list[str] = []

        # Add an overall title when requested
        if doc_title:
            html_parts.append(f"<h1>{doc_title}</h1>")

        # Group the DataFrame and render each subgroup
        grouped = df.groupby(group_cols, dropna=False)

        for keys, group in grouped:
            # Ensure *keys* is always iterable
            keys = (keys,) if not isinstance(keys, tuple) else keys
            group_name = ", ".join(f"{col}: {val}" for col, val in zip(group_cols, keys))

            html_parts.append(f"<h2>{group_name}</h2>")
            html_parts.append(
                group.to_html(index=False, escape=False, border=0, classes="dataframe")
            )

        return "\n".join(html_parts)


    # 🗂  USER HISTORY
    # Expander elements have been replaced with headers to avoid nesting issues
    try:
        with st.expander("📅 View your previously uploaded CSV files", expanded=False):
            user_history = get_user_history(user_id)
            if user_history:
                st.write("Previously uploaded CSV files (click to load):")
                for idx, entry in enumerate(user_history):
                    if isinstance(entry, dict) and 'file_name' in entry and 'timestamp' in entry:
                        # If entry has cache_id, make it clickable
                        if 'cache_id' in entry:
                            if st.button(f"Load: {entry['file_name']}", key=f"history_{idx}"):
                                # Load the cached CSV and options
                                df, options = get_cached_csv(entry['cache_id'])
                                if df is not None:
                                    # Add to dataframes list if not already there
                                    if not any(df.equals(existing_df) for existing_df in st.session_state.dataframes):
                                        # Store the dataframe
                                        st.session_state.dataframes.append(df)

                                        # Store the options for this dataframe if available
                                        if options:
                                            if "dataframe_options" not in st.session_state:
                                                st.session_state.dataframe_options = {}
                                            # Use the length of dataframes list - 1 as the index for this dataframe
                                            df_idx = len(st.session_state.dataframes) - 1
                                            st.session_state.dataframe_options[df_idx] = options

                                        st.success(f"Loaded {entry['file_name']} from history")
                                        st.rerun()
                                    else:
                                        st.info(f"{entry['file_name']} is already loaded")
                                else:
                                    st.error(f"Could not load {entry['file_name']} from cache")
                            # Display timestamp separately
                            st.caption(f"Uploaded: {entry['timestamp']}")
                        else:
                            # For entries without cache_id, just display them
                            st.write(f"{idx+1}. **{entry['file_name']}** - {entry['timestamp']}")
                    elif isinstance(entry, str):
                        # Handle legacy entries that might only contain filenames
                        st.write(f"{idx+1}. **{entry}**")
                    else:
                        st.write(f"{idx+1}. **Unknown entry format**")
            else:
                st.write("No CSV files have been uploaded yet.")
    except Exception as e:
        st.error(f"Error loading user history: {e}")
        st.markdown(f"**Exception:**<br>{traceback.format_exc().replace('\n', '<br>')}", unsafe_allow_html=True)

    # Display current session cache
    st.header("Current Session Cache")
    try:
        cached_csvs = get_all_cached_csvs()
        if cached_csvs:
            st.write("CSVs cached in current session:")
            for cache_id, cache_info in cached_csvs.items():
                # Create a button to reload this CSV
                if st.button(f"Load: {cache_info['file_name']}", key=f"load_{cache_id}"):
                    # Add this cached CSV to dataframes for processing
                    df = cache_info['df']
                    options = cache_info.get('options')
                    if df is not None and not any(df.equals(existing_df) for existing_df in st.session_state.dataframes):
                        # Store the dataframe
                        st.session_state.dataframes.append(df)

                        # Store the options for this dataframe if available
                        if options:
                            if "dataframe_options" not in st.session_state:
                                st.session_state.dataframe_options = {}
                            # Use the length of dataframes list - 1 as the index for this dataframe
                            df_idx = len(st.session_state.dataframes) - 1
                            st.session_state.dataframe_options[df_idx] = options

                        st.success(f"Loaded cached CSV: {cache_info['file_name']}")
                        st.experimental_rerun()
        else:
            st.write("No CSVs are cached in the current session.")
    except Exception as e:
        st.error(f"Error handling session cache: {e}")
        st.markdown(f"**Exception:**<br>{traceback.format_exc().replace('\n', '<br>')}", unsafe_allow_html=True)
